from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_report():
    doc = Document()

    # Title
    title = doc.add_heading('PhonePe Pulse Data Visualization and Analysis', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Author
    p = doc.add_paragraph()
    run = p.add_run('Author: Nikhil Kumar')
    run.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph()
    run = p.add_run('LabMentix Data Science Internship Project')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Section: Executive Summary
    doc.add_heading('Executive Summary', level=1)
    doc.add_paragraph(
        "This project aims to analyze and visualize the PhonePe Pulse data from 2018 to 2024. "
        "The objective is to provide an interactive dashboard that allows users to explore transaction trends, "
        "user demographics, and insurance penetration across India. The project also includes a machine learning "
        "component to forecast future transaction amounts."
    )

    # Section: Technologies Used
    doc.add_heading('Technologies Used', level=1)
    techs = [
        "Python (Core Programming)",
        "Streamlit (Dashboard Framework)",
        "Plotly (Interactive Visualizations)",
        "SQLite (Data Storage)",
        "Scikit-learn (Machine Learning)",
        "Pandas (Data Manipulation)"
    ]
    for tech in techs:
        doc.add_paragraph(tech, style='List Bullet')

    # Section: Problem Statement
    doc.add_heading('Problem Statement', level=1)
    doc.add_paragraph(
        "The PhonePe Pulse repository contains vast amounts of JSON data. Manually extracting insights "
        "from these files is nearly impossible for non-technical users. This project solves that by "
        "building a comprehensive ETL pipeline and a user-friendly web interface."
    )

    # Section: ETL Pipeline
    doc.add_heading('Data Processing (ETL)', level=1)
    doc.add_paragraph(
        "The ETL pipeline (scripts/data_etl.py) performs the following steps:"
    )
    steps = [
        "Extraction: Cloning the PhonePe Pulse GitHub repository and reading raw JSON files.",
        "Transformation: Cleaning and structuring data into a tabular format using Pandas.",
        "Loading: Storing the processed data into a local SQLite database (phonepe_pulse.db) for fast querying."
    ]
    for step in steps:
        doc.add_paragraph(step, style='List Number')

    # Section: Features
    doc.add_heading('Core Features', level=1)
    features = [
        "Geographical Analysis: Choropleth maps visualizing transaction hotspots.",
        "Transaction Insights: Yearly and quarterly trends split by categories.",
        "Insurance Analysis: Deep dive into digital insurance growth.",
        "Top Rankings: Leaderboards for states and districts.",
        "Growth Predictions: A Random Forest Regressor to predict future quarterly amounts."
    ]
    for feature in features:
        doc.add_paragraph(feature, style='List Bullet')

    # Section: Machine Learning
    doc.add_heading('Machine Learning Component', level=1)
    doc.add_paragraph(
        "The project uses a Random Forest Regressor trained on historical quarterly data. "
        "The model predicts the 'Transaction Amount' based on 'State', 'Year', 'Quarter', and 'Transaction Type'. "
        "Accuracy is optimized through feature encoding and rigorous testing."
    )

    # Section: Deployment
    doc.add_heading('Deployment', level=1)
    doc.add_paragraph(
        "The application is deployed on Streamlit Community Cloud for global accessibility."
    )
    p = doc.add_paragraph("Live URL: ")
    p.add_run("https://viraj357-phonepay-app-jokpqy.streamlit.app/").bold = True

    # Save
    doc.save('PhonePe_Pulse_Report.docx')
    print("Report generated: PhonePe_Pulse_Report.docx")

if __name__ == "__main__":
    create_report()
